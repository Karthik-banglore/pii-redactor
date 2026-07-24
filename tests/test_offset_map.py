"""Critical offset-mapping tests — the correctness of the whole tool hangs here."""

from __future__ import annotations

from docx import Document
from docx.shared import Pt

from pii_redactor.adapters.docx.segments import ParagraphView
from pii_redactor.domain.span import ResolvedSpan


def _fragmented_paragraph(words):
    """Build a paragraph with one run per word (and space runs between)."""
    doc = Document()
    para = doc.add_paragraph()
    for i, word in enumerate(words):
        if i > 0:
            para.add_run(" ")
        run = para.add_run(word)
        run.bold = i == 0  # first word bold — formatting preservation check
        run.font.size = Pt(12)
    return doc, para


def _rs(start, end, original, replacement, etype="PERSON"):
    return ResolvedSpan(
        start=start,
        end=end,
        entity_type=etype,
        original=original,
        replacement=replacement,
        score=1.0,
        source="test",
    )


def test_single_run_replacement():
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("Hello world")
    view = ParagraphView(para)
    assert view.text == "Hello world"
    view.apply([_rs(0, 5, "Hello", "Hi")])
    assert ParagraphView(para).text == "Hi world"


def test_five_run_name_replacement():
    """Rajesh Kushal Hegde split across five runs (word/space/word/space/word)."""
    doc, para = _fragmented_paragraph(["Rajesh", "Kushal", "Hegde"])
    view = ParagraphView(para)
    assert view.text == "Rajesh Kushal Hegde"
    assert len(view.run_map) == 5  # 3 words + 2 spaces

    view.apply([_rs(0, 19, "Rajesh Kushal Hegde", "Arjun Mehta")])
    rebuilt = ParagraphView(para)
    assert rebuilt.text == "Arjun Mehta"
    # First run keeps bold formatting
    assert para.runs[0].bold is True


def test_mid_run_start_and_end():
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("Contact rajesh@ksh.com today")
    view = ParagraphView(para)
    # "rajesh@ksh.com" at [8, 22)
    view.apply([_rs(8, 22, "rajesh@ksh.com", "alice@example.com", "EMAIL")])
    assert ParagraphView(para).text == "Contact alice@example.com today"


def test_two_spans_right_to_left():
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("Alice and Bob are here")
    view = ParagraphView(para)
    # Alice [0,5), Bob [10,13)
    view.apply(
        [
            _rs(0, 5, "Alice", "Ava"),
            _rs(10, 13, "Bob", "Ben"),
        ]
    )
    assert ParagraphView(para).text == "Ava and Ben are here"


def test_shorter_and_longer_replacement():
    doc, para = _fragmented_paragraph(["Rajesh", "Kushal", "Hegde"])
    view = ParagraphView(para)
    view.apply([_rs(0, 19, "Rajesh Kushal Hegde", "Jo")])
    assert ParagraphView(para).text == "Jo"

    doc2, para2 = _fragmented_paragraph(["Jo"])
    view2 = ParagraphView(para2)
    view2.apply([_rs(0, 2, "Jo", "Jonathan Alexander")])
    assert ParagraphView(para2).text == "Jonathan Alexander"


def test_formatting_preserved_on_first_run():
    doc, para = _fragmented_paragraph(["Bold", "Name", "Here"])
    assert para.runs[0].bold is True
    view = ParagraphView(para)
    view.apply([_rs(0, 14, "Bold Name Here", "Replacement")])
    assert para.runs[0].bold is True
    assert para.runs[0].text.startswith("Replacement")


def test_empty_paragraph_noop():
    doc = Document()
    para = doc.add_paragraph()
    view = ParagraphView(para)
    assert view.text == ""
    view.apply([])
    assert ParagraphView(para).text == ""


def test_identity_roundtrip_preserves_text():
    """Applying no spans must leave text unchanged."""
    doc, para = _fragmented_paragraph(["Kushal", "Subbayya", "Hegde"])
    before = ParagraphView(para).text
    ParagraphView(para).apply([])
    assert ParagraphView(para).text == before
